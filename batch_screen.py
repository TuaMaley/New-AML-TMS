"""
Batch Sanctions Screening — File Parser + Report Generator
============================================================
Accepts Word (.docx), Excel (.xlsx/.xls), or PDF files containing
a list of names, screens each against the SDN list, and produces
a formatted Excel report for download.
"""
import io, re, os
from datetime import datetime

# ── Optional dependency checks ────────────────────────────────────────────────
try:
    from docx import Document as _DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import openpyxl as _openpyxl
    _HAS_OPENPYXL = True
except ImportError:
    _HAS_OPENPYXL = False

try:
    from pypdf import PdfReader as _PdfReader
    _HAS_PYPDF = True
except ImportError:
    _HAS_PYPDF = False


# ── Name extraction ───────────────────────────────────────────────────────────

def extract_names_from_docx(file_bytes: bytes) -> list:
    if not _HAS_DOCX:
        raise ValueError("python-docx not installed. Run: pip install python-docx")
    doc = _DocxDocument(io.BytesIO(file_bytes))
    names = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            for part in re.split(r'[\n;|,]', text):
                part = part.strip()
                if len(part) >= 2:
                    names.append(part)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if len(text) >= 2 and text not in names:
                    names.append(text)
    return names


def extract_names_from_xlsx(file_bytes: bytes) -> list:
    if not _HAS_OPENPYXL:
        raise ValueError("openpyxl not installed. Run: pip install openpyxl")
    wb = _openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    names = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            for cell in row:
                if cell and isinstance(cell, str):
                    text = cell.strip()
                    if len(text) >= 2:
                        names.append(text)
    wb.close()
    return names


def extract_names_from_pdf(file_bytes: bytes) -> list:
    if not _HAS_PYPDF:
        raise ValueError("pypdf not installed. Run: pip install pypdf")
    reader = _PdfReader(io.BytesIO(file_bytes))
    names = []
    for page in reader.pages:
        text = page.extract_text() or ''
        for line in text.split('\n'):
            line = line.strip()
            if len(line) >= 2:
                for part in re.split(r'[;|]', line):
                    part = part.strip()
                    if len(part) >= 2:
                        names.append(part)
    return names


def extract_names(file_bytes: bytes, filename: str) -> list:
    """Dispatch to correct parser based on file extension."""
    ext = os.path.splitext(filename.lower())[1]
    try:
        if ext in ('.docx', '.doc'):
            raw = extract_names_from_docx(file_bytes)
        elif ext in ('.xlsx', '.xls', '.xlsm'):
            raw = extract_names_from_xlsx(file_bytes)
        elif ext == '.pdf':
            raw = extract_names_from_pdf(file_bytes)
        else:
            text = file_bytes.decode('utf-8', errors='replace')
            raw = [l.strip() for l in text.splitlines() if len(l.strip()) >= 2]
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Could not parse file '{filename}': {e}")

    skip_patterns = re.compile(
        r'^(name|entity|customer|client|company|first|last|full|#|no\.|sr\.?|s/n|sl\.?)',
        re.IGNORECASE
    )
    seen = set()
    names = []
    for n in raw:
        key = n.lower().strip()
        if key not in seen and not skip_patterns.match(n) and len(n) >= 2:
            seen.add(key)
            names.append(n)
    return names[:500]


# ── Report generator ──────────────────────────────────────────────────────────

def generate_report(results: list, filename: str) -> bytes:
    """Generate a professional Excel report. Returns .xlsx bytes."""
    if not _HAS_OPENPYXL:
        raise ValueError("openpyxl not installed. Run: pip install openpyxl")

    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = _openpyxl.Workbook()

    # ── Colors ──────────────────────────────────────────────────────────────
    NAVY      = "0F285A"
    RED       = "C0392B"
    ORANGE    = "E67E22"
    GREEN     = "1E8449"
    YELLOW    = "F39C12"
    LIGHT_RED = "FADBD8"
    LIGHT_ORG = "FAD7A0"
    LIGHT_GRN = "D5F5E3"
    LIGHT_YEL = "FEF9E7"
    LIGHT_GRY = "F2F3F4"
    WHITE     = "FFFFFF"

    def fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def bdr():
        thin = Side(style='thin', color="CCCCCC")
        return Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── Summary sheet ────────────────────────────────────────────────────────
    ws_sum = wb.active
    ws_sum.title = "Summary"
    ws_sum.sheet_view.showGridLines = False
    ws_sum.column_dimensions['A'].width = 35
    ws_sum.column_dimensions['B'].width = 20

    ws_sum.merge_cells('A1:B1')
    c = ws_sum['A1']
    c.value = "AML-TMS SANCTIONS SCREENING REPORT"
    c.font  = Font(name='Calibri', bold=True, size=16, color=WHITE)
    c.fill  = fill(NAVY)
    c.alignment = Alignment(horizontal='center', vertical='center')
    ws_sum.row_dimensions[1].height = 36

    ws_sum.merge_cells('A2:B2')
    c2 = ws_sum['A2']
    c2.value = f"Source: {filename}   |   Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    c2.font  = Font(name='Calibri', size=10, color=WHITE, italic=True)
    c2.fill  = fill("1E3A5F")
    c2.alignment = Alignment(horizontal='center', vertical='center')
    ws_sum.row_dimensions[2].height = 20

    total     = len(results)
    confirmed = sum(1 for r in results if r['status'] == 'HIT')
    potential = sum(1 for r in results if r['status'] == 'POTENTIAL_MATCH')
    weak      = sum(1 for r in results if r['status'] == 'WEAK_MATCH')
    clear     = sum(1 for r in results if r['status'] == 'CLEAR')
    action    = confirmed + potential

    stats = [
        ("Total names screened",      total,     NAVY,   WHITE),
        ("CONFIRMED HITS (>=80%)",    confirmed, RED,    WHITE),
        ("POTENTIAL MATCHES (>=70%)", potential, ORANGE, WHITE),
        ("WEAK MATCHES (55-69%)",     weak,      YELLOW, NAVY),
        ("CLEAR — no match",          clear,     GREEN,  WHITE),
        ("Action required",           action,    RED if action else GREEN, WHITE),
    ]

    row = 4
    for label, value, bg, fg in stats:
        ws_sum[f'A{row}'].value = label
        ws_sum[f'A{row}'].font  = Font(name='Calibri', size=11, color=NAVY)
        ws_sum[f'A{row}'].fill  = fill(LIGHT_GRY)
        ws_sum[f'A{row}'].border = bdr()
        ws_sum[f'A{row}'].alignment = Alignment(vertical='center', indent=1)
        ws_sum.row_dimensions[row].height = 22
        ws_sum[f'B{row}'].value = value
        ws_sum[f'B{row}'].font  = Font(name='Calibri', bold=True, size=13, color=fg)
        ws_sum[f'B{row}'].fill  = fill(bg)
        ws_sum[f'B{row}'].border = bdr()
        ws_sum[f'B{row}'].alignment = Alignment(horizontal='center', vertical='center')
        row += 1

    row += 1
    ws_sum[f'A{row}'].value = "SDN List version"
    ws_sum[f'A{row}'].font  = Font(name='Calibri', size=10, italic=True, color="888888")
    ws_sum[f'B{row}'].value = "OFAC SDN Apr 3 2026 (sdn_advanced.xml) — 17,007 entries"
    ws_sum[f'B{row}'].font  = Font(name='Calibri', size=10, italic=True, color="888888")

    # ── Results sheet ────────────────────────────────────────────────────────
    ws = wb.create_sheet("Screening Results")
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = 'A2'

    col_widths = [5, 30, 20, 18, 18, 12, 35, 30, 18, 18]
    col_names  = ['#', 'Queried Name', 'Status', 'Risk Level',
                  'Confirmed Hits', 'Score', 'SDN Match Name',
                  'Sanction Programs', 'Country', 'Date Added']
    for i, (w, n) in enumerate(zip(col_widths, col_names), 1):
        col = get_column_letter(i)
        ws.column_dimensions[col].width = w
        c = ws.cell(row=1, column=i, value=n)
        c.font      = Font(name='Calibri', bold=True, size=11, color=WHITE)
        c.fill      = fill(NAVY)
        c.border    = bdr()
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    ws.row_dimensions[1].height = 28

    status_styles = {
        'HIT':             (fill(LIGHT_RED), RED,    'CONFIRMED HIT'),
        'POTENTIAL_MATCH': (fill(LIGHT_ORG), ORANGE, 'POTENTIAL MATCH'),
        'WEAK_MATCH':      (fill(LIGHT_YEL), YELLOW, 'WEAK MATCH'),
        'CLEAR':           (fill(LIGHT_GRN), GREEN,  'CLEAR'),
        'error':           (fill(LIGHT_GRY), "888888",'ERROR'),
    }

    for idx, r in enumerate(results, 1):
        row_n = idx + 1
        st    = r.get('status', 'error')
        bg_fill, fg_color, st_label = status_styles.get(st, status_styles['error'])
        matches  = r.get('matches', [])
        best     = matches[0] if matches else {}
        sdn_name = re.sub(r'[\u0400-\u04FF]', '', best.get('sdn_name', '—')).strip()
        programs = ', '.join(best.get('programs', []))[:60] if best else '—'
        country  = best.get('country', '—') if best else '—'
        date_add = best.get('date_added', '—') if best else '—'
        score    = best.get('match_score', '—') if best else '—'
        confirmed_hits = r.get('confirmed_hits', 0)

        row_data = [
            idx, r.get('query', ''), st_label,
            r.get('risk_level', '').upper(), confirmed_hits,
            f"{score}%" if isinstance(score, int) else '—',
            sdn_name, programs, country, date_add,
        ]
        for col_i, val in enumerate(row_data, 1):
            c = ws.cell(row=row_n, column=col_i, value=val)
            c.font   = Font(name='Calibri', size=10,
                            color=fg_color if col_i in (3, 4) else NAVY,
                            bold=(col_i in (3, 4)))
            c.fill   = bg_fill
            c.border = bdr()
            c.alignment = Alignment(vertical='center',
                                    wrap_text=(col_i in (7, 8)),
                                    horizontal='center' if col_i in (1, 5, 6) else 'left')
        ws.row_dimensions[row_n].height = 18

    # ── Hits detail sheet ────────────────────────────────────────────────────
    hits = [r for r in results if r['status'] in ('HIT', 'POTENTIAL_MATCH')]
    if hits:
        ws_h = wb.create_sheet("Hits & Matches Detail")
        ws_h.sheet_view.showGridLines = False
        ws_h.freeze_panes = 'A2'

        detail_cols  = [5, 28, 16, 26, 26, 12, 28, 55, 18, 18]
        detail_heads = ['#', 'Queried Name', 'Status', 'SDN Name', 'Matched On',
                        'Score', 'Programs', 'Reason for Designation', 'Country', 'Date Added']
        for i, (w, n) in enumerate(zip(detail_cols, detail_heads), 1):
            col = get_column_letter(i)
            ws_h.column_dimensions[col].width = w
            c = ws_h.cell(row=1, column=i, value=n)
            c.font      = Font(name='Calibri', bold=True, size=11, color=WHITE)
            c.fill      = fill(RED)
            c.border    = bdr()
            c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        ws_h.row_dimensions[1].height = 28

        detail_row = 2
        for r in hits:
            for m in (r.get('matches') or [{}])[:3]:
                st    = r.get('status', '')
                bg_f  = fill(LIGHT_RED) if st == 'HIT' else fill(LIGHT_ORG)
                fg_c  = RED if st == 'HIT' else ORANGE
                sdn_n = re.sub(r'[\u0400-\u04FF]', '', m.get('sdn_name', '—')).strip()
                reason = (m.get('reason', '—') or '—')[:120]
                row_data = [
                    detail_row - 1, r.get('query', ''),
                    'CONFIRMED HIT' if st == 'HIT' else 'POTENTIAL MATCH',
                    sdn_n, m.get('match_on', '—'),
                    f"{m.get('match_score', '—')}%" if m.get('match_score') else '—',
                    ', '.join(m.get('programs', [])),
                    reason, m.get('country', '—'), m.get('date_added', '—'),
                ]
                for col_i, val in enumerate(row_data, 1):
                    c = ws_h.cell(row=detail_row, column=col_i, value=val)
                    c.font   = Font(name='Calibri', size=10,
                                    color=fg_c if col_i == 3 else NAVY,
                                    bold=(col_i == 3))
                    c.fill   = bg_f
                    c.border = bdr()
                    c.alignment = Alignment(vertical='center',
                                            wrap_text=(col_i == 8),
                                            horizontal='center' if col_i in (1, 6) else 'left')
                ws_h.row_dimensions[detail_row].height = 18
                detail_row += 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
